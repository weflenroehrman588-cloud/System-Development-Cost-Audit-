import importlib.util
import json
import tempfile
import unittest
from copy import deepcopy
from decimal import Decimal
from contextlib import redirect_stderr
from io import StringIO
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "generate_detailed_quotation.py"
TEMPLATE = ROOT / "templates" / "detailed-quotation-input.json"
SPEC = importlib.util.spec_from_file_location("generate_detailed_quotation", SCRIPT)
quotation = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(quotation)


class DetailedQuotationTest(unittest.TestCase):
    def setUp(self):
        self.data = json.loads(TEMPLATE.read_text(encoding="utf-8"))

    def test_template_totals_and_reconciliation(self):
        result = quotation.prepare(self.data)
        self.assertEqual(result["total_days"], Decimal("45"))
        self.assertEqual(result["day_rate"], Decimal("1200.55"))
        self.assertEqual(result["base_amount"], Decimal("54024.75"))
        self.assertEqual(result["risk_amount"], Decimal("8103.71"))
        self.assertEqual(result["total_amount"], Decimal("62128.46"))
        self.assertEqual(
            sum((row["days"] for row in result["people"]), Decimal("0")),
            result["total_days"],
        )
        self.assertEqual(
            sum((row["base_amount"] for row in result["people"]), Decimal("0")),
            result["base_amount"],
        )
        self.assertEqual(
            sum((row["risk_amount"] for row in result["modules"]), Decimal("0")),
            result["risk_amount"],
        )
        self.assertEqual(
            sum((row["payment_amount"] for row in result["implementation_plan"]), Decimal("0")),
            result["total_amount"],
        )
        self.assertTrue(all(row["base_amount"] >= 0 for row in result["people"]))
        self.assertTrue(all(row["risk_amount"] >= 0 for row in result["modules"]))

    def test_rejects_item_person_day_mismatch(self):
        data = deepcopy(self.data)
        data["modules"][0]["items"][0]["person_days"]["协调A"] = 1
        with self.assertRaisesRegex(ValueError, "不等于功能工作量"):
            quotation.prepare(data)

    def test_rejects_unknown_person(self):
        data = deepcopy(self.data)
        allocation = data["modules"][0]["items"][0]["person_days"]
        allocation["未登记人员"] = allocation.pop("协调A")
        with self.assertRaisesRegex(ValueError, "未登记人员"):
            quotation.prepare(data)

    def test_largest_remainder_allocation_never_goes_negative(self):
        allocation = quotation.allocate_amount_by_weight(
            Decimal("0.03"),
            [(f"人员{index}", Decimal("1")) for index in range(5)],
            Decimal("5"),
        )
        self.assertEqual(sum(allocation.values(), Decimal("0")), Decimal("0.03"))
        self.assertTrue(all(value >= 0 for value in allocation.values()))

    def test_rejects_more_than_two_decimal_places_for_days(self):
        data = deepcopy(self.data)
        item = data["modules"][0]["items"][0]
        item["days"] = 3.005
        item["person_days"] = {"负责人A": 1.505, "协调A": 1.5}
        with self.assertRaisesRegex(ValueError, "最多保留两位小数"):
            quotation.prepare(data)

    def test_rejects_payment_ratio_not_equal_to_one_hundred_percent(self):
        data = deepcopy(self.data)
        data["implementation_plan"][-1]["payment_ratio"] = "20%"
        with self.assertRaisesRegex(ValueError, "合计必须为 100%"):
            quotation.prepare(data)

    def test_json_preserves_ratio_precision(self):
        data = deepcopy(self.data)
        ratios = ("33.33%", "33.33%", "33.34%")
        for row, ratio in zip(data["implementation_plan"], ratios):
            row["payment_ratio"] = ratio
        payload = quotation.serializable(quotation.prepare(data))
        self.assertEqual(
            sum((Decimal(row["payment_ratio"]) for row in payload["implementation_plan"]), Decimal("0")),
            Decimal("1"),
        )
        self.assertEqual(payload["implementation_plan"][-1]["payment_ratio"], "0.3334")
        self.assertEqual(payload["risk_rate"], "0.15")

    def test_enforces_cny_and_evidence_grade_domain(self):
        currency_data = deepcopy(self.data)
        currency_data["meta"]["currency"] = "USD"
        with self.assertRaisesRegex(ValueError, "必须为 CNY"):
            quotation.prepare(currency_data)
        evidence_data = deepcopy(self.data)
        evidence_data["modules"][0]["items"][0]["evidence_grade"] = "Z"
        with self.assertRaisesRegex(ValueError, "必须是 A 至 E"):
            quotation.prepare(evidence_data)
        normalized_data = deepcopy(self.data)
        normalized_data["meta"]["currency"] = "cny"
        normalized_data["modules"][0]["items"][0]["evidence_grade"] = "b"
        result = quotation.prepare(normalized_data)
        self.assertEqual(result["meta"]["currency"], "CNY")
        self.assertEqual(result["items"][0]["evidence_grade"], "B")

    def test_cli_rejects_non_object_and_non_finite_number_without_traceback(self):
        for content in ("[]", '{"meta":{"monthly_rate_cny":NaN}}'):
            with self.subTest(content=content), tempfile.TemporaryDirectory() as directory:
                source = Path(directory) / "invalid.json"
                source.write_text(content, encoding="utf-8")
                stderr = StringIO()
                with redirect_stderr(stderr):
                    exit_code = quotation.main([str(source), "--formats", "json"])
                self.assertEqual(exit_code, 2)
                self.assertIn("错误：", stderr.getvalue())
                self.assertNotIn("Traceback", stderr.getvalue())

    def test_generates_core_outputs(self):
        with tempfile.TemporaryDirectory() as directory:
            exit_code = quotation.main(
                [
                    str(TEMPLATE),
                    "--output-dir",
                    directory,
                    "--prefix",
                    "示例",
                    "--formats",
                    "json,csv,md",
                ]
            )
            self.assertEqual(exit_code, 0)
            output_dir = Path(directory)
            calculation = json.loads(
                (output_dir / "示例-calculation.json").read_text(encoding="utf-8")
            )
            self.assertEqual(calculation["total_amount"], "62128.46")
            self.assertIn("current_system_baseline", calculation)
            self.assertIn("implementation_plan", calculation)
            self.assertIn("exclusions", calculation)
            markdown = (output_dir / "示例-quotation.md").read_text(encoding="utf-8")
            for heading in ("分部分项工作量与报价", "现有系统功能基线", "模块—人员工作量矩阵", "实施与付款计划"):
                self.assertIn(heading, markdown)
            self.assertGreater((output_dir / "示例-details.csv").stat().st_size, 100)

    def test_spreadsheet_text_is_escaped(self):
        data = deepcopy(self.data)
        data["modules"][0]["items"][0]["name"] = "=HYPERLINK(\"https://invalid.example\")"
        result = quotation.prepare(data)
        with tempfile.TemporaryDirectory() as directory:
            csv_path = Path(directory) / "safe.csv"
            quotation.write_csv(result, csv_path)
            self.assertIn("'=HYPERLINK", csv_path.read_text(encoding="utf-8-sig"))
            try:
                import openpyxl
            except ImportError:
                return
            xlsx_path = Path(directory) / "safe.xlsx"
            quotation.write_xlsx(result, xlsx_path)
            workbook = openpyxl.load_workbook(xlsx_path, data_only=False)
            cell = workbook["分部分项清单"]["D2"]
            self.assertEqual(cell.data_type, "s")
            self.assertTrue(cell.value.startswith("'=HYPERLINK"))

    def test_preflight_failure_leaves_no_partial_outputs(self):
        original = quotation._docx_imports
        quotation._docx_imports = lambda: (_ for _ in ()).throw(RuntimeError("missing"))
        try:
            with tempfile.TemporaryDirectory() as directory:
                stderr = StringIO()
                with redirect_stderr(stderr):
                    exit_code = quotation.main([
                        str(TEMPLATE), "--output-dir", directory, "--prefix", "半套", "--formats", "json,docx"
                    ])
                self.assertEqual(exit_code, 2)
                self.assertEqual(list(Path(directory).iterdir()), [])
        finally:
            quotation._docx_imports = original

    def test_publish_failure_rolls_back_all_new_outputs(self):
        original = quotation.os.replace
        calls = 0

        def fail_second_replace(source, target):
            nonlocal calls
            calls += 1
            if calls == 2:
                raise OSError("simulated publish failure")
            return original(source, target)

        quotation.os.replace = fail_second_replace
        try:
            with tempfile.TemporaryDirectory() as directory:
                stderr = StringIO()
                with redirect_stderr(stderr):
                    exit_code = quotation.main([
                        str(TEMPLATE), "--output-dir", directory, "--prefix", "回滚", "--formats", "json,csv"
                    ])
                self.assertEqual(exit_code, 2)
                self.assertEqual(list(Path(directory).iterdir()), [])
                self.assertIn("已回滚", stderr.getvalue())
        finally:
            quotation.os.replace = original

    def test_generates_office_outputs_when_dependencies_exist(self):
        try:
            import openpyxl  # noqa: F401
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("未安装 Office 文档可选依赖")
        with tempfile.TemporaryDirectory() as directory:
            result = quotation.prepare(self.data)
            xlsx = Path(directory) / "quotation.xlsx"
            docx_path = Path(directory) / "quotation.docx"
            quotation.write_xlsx(result, xlsx)
            quotation.write_docx(result, docx_path)
            self.assertTrue(xlsx.exists())
            self.assertTrue(docx_path.exists())
            workbook = openpyxl.load_workbook(xlsx, data_only=False)
            self.assertIn("分部分项清单", workbook.sheetnames)
            self.assertTrue(workbook["分部分项清单"]["J2"].value.startswith("=ROUND("))
            from docx import Document

            document = Document(docx_path)
            section = document.sections[0]
            self.assertLess(section.page_width, section.page_height)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            for heading in ("现有系统功能基线", "模块—人员工作量矩阵", "实施与付款计划"):
                self.assertIn(heading, text)
            self.assertGreaterEqual(len(document.tables), 10)


if __name__ == "__main__":
    unittest.main()
