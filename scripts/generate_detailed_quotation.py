#!/usr/bin/env python3
"""从结构化 WBS 生成可复算的工作量统计与报价成果。

核心格式（JSON/Markdown/CSV）仅依赖 Python 标准库；XLSX 和 DOCX 分别需要
openpyxl 与 python-docx。所有金额采用 Decimal 和 ROUND_HALF_UP。
"""

from __future__ import annotations

import argparse
import csv
import json
import os
import sys
import tempfile
from collections import defaultdict
from decimal import Decimal, ROUND_FLOOR, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Iterable


MONEY = Decimal("0.01")
DAY = Decimal("0.01")


def decimal(value: Any, field: str) -> Decimal:
    try:
        result = Decimal(str(value))
    except Exception as exc:  # pragma: no cover - Decimal 异常类型较多
        raise ValueError(f"{field} 必须是数值，当前为 {value!r}") from exc
    if not result.is_finite():
        raise ValueError(f"{field} 必须是有限数，当前为 {value!r}")
    return result


def round_money(value: Decimal) -> Decimal:
    return value.quantize(MONEY, rounding=ROUND_HALF_UP)


def round_day(value: Decimal) -> Decimal:
    return value.quantize(DAY, rounding=ROUND_HALF_UP)


def money_text(value: Decimal) -> str:
    return f"{value:.2f}"


def number_text(value: Decimal) -> str:
    normalized = value.normalize()
    return format(normalized, "f")


def require_text(mapping: dict[str, Any], key: str, context: str) -> str:
    value = str(mapping.get(key, "")).strip()
    if not value:
        raise ValueError(f"{context}.{key} 不能为空")
    return value


def allocate_amount_by_weight(
    amount: Decimal, allocations: list[tuple[str, Decimal]], total_weight: Decimal
) -> dict[str, Decimal]:
    """用最大余数法分配整数分，保证金额非负且合计严格相等。"""
    if amount < 0 or total_weight < 0:
        raise ValueError("待分摊金额和权重合计不能为负数")
    if not allocations:
        if amount:
            raise ValueError("非零金额缺少分摊对象")
        return {}
    if any(weight < 0 for _, weight in allocations):
        raise ValueError("分摊权重不能为负数")
    if total_weight == 0:
        if amount:
            raise ValueError("非零金额不能按零权重分摊")
        return {name: Decimal("0.00") for name, _ in allocations}

    total_cents = int((round_money(amount) * 100).to_integral_exact())
    floors: list[int] = []
    remainders: list[tuple[Decimal, int]] = []
    for index, (_, weight) in enumerate(allocations):
        ideal = Decimal(total_cents) * weight / total_weight
        floor_cents = int(ideal.to_integral_value(rounding=ROUND_FLOOR))
        floors.append(floor_cents)
        remainders.append((ideal - floor_cents, index))
    remaining = total_cents - sum(floors)
    order = [index for _, index in sorted(remainders, key=lambda item: (-item[0], item[1]))]
    for index in order[:remaining]:
        floors[index] += 1
    return {
        name: Decimal(cents) / Decimal("100")
        for (name, _), cents in zip(allocations, floors)
    }


def spreadsheet_safe(value: Any) -> str:
    """阻止外部文本被 Excel/CSV 解释为公式。"""
    text = str(value)
    if text.lstrip().startswith(("=", "+", "-", "@", "\t", "\r")):
        return "'" + text
    return text


def parse_payment_ratio(value: Any, field: str) -> Decimal:
    if isinstance(value, str) and value.strip().endswith("%"):
        ratio = decimal(value.strip()[:-1], field) / Decimal("100")
    else:
        ratio = decimal(value, field)
    if ratio < 0 or ratio > 1:
        raise ValueError(f"{field} 必须在 0% 到 100% 之间")
    return ratio


def prepare(data: dict[str, Any]) -> dict[str, Any]:
    """校验输入并生成明细、模块、人员及金额勾稽结果。"""
    if not isinstance(data, dict):
        raise ValueError("输入 JSON 顶层必须是对象")
    meta_input = data.get("meta")
    if not isinstance(meta_input, dict):
        raise ValueError("meta 必须是对象")
    meta = dict(meta_input)
    project_name = require_text(meta, "project_name", "meta")
    for key in ("price_base_date", "currency", "tax_basis"):
        require_text(meta, key, "meta")
    if str(meta["currency"]).strip().upper() != "CNY":
        raise ValueError("meta.currency 必须为 CNY；其他币种需另行定义汇率和金额单位")
    meta["currency"] = "CNY"
    monthly_rate = decimal(meta.get("monthly_rate_cny"), "meta.monthly_rate_cny")
    month_days = decimal(meta.get("person_month_days"), "meta.person_month_days")
    risk_rate = decimal(meta.get("risk_reserve_rate", 0), "meta.risk_reserve_rate")
    if monthly_rate <= 0 or month_days <= 0:
        raise ValueError("月费率和每人月人日必须大于 0")
    if risk_rate < 0:
        raise ValueError("风险准备率不能小于 0")
    day_rate = round_money(monthly_rate / month_days)

    personnel = data.get("personnel")
    if not isinstance(personnel, list) or not personnel:
        raise ValueError("personnel 必须是非空数组")
    people: dict[str, dict[str, Any]] = {}
    for index, person in enumerate(personnel, 1):
        if not isinstance(person, dict):
            raise ValueError(f"personnel[{index}] 必须是对象")
        name = require_text(person, "name", f"personnel[{index}]")
        require_text(person, "role", f"personnel[{index}]")
        require_text(person, "responsibility", f"personnel[{index}]")
        if name in people:
            raise ValueError(f"人员姓名重复：{name}")
        people[name] = person

    baseline_input = data.get("current_system_baseline")
    if not isinstance(baseline_input, list):
        raise ValueError("current_system_baseline 必须是数组")
    baseline: list[dict[str, str]] = []
    for index, row in enumerate(baseline_input, 1):
        if not isinstance(row, dict):
            raise ValueError(f"current_system_baseline[{index}] 必须是对象")
        context = f"current_system_baseline[{index}]"
        baseline.append(
            {
                "id": str(row.get("id", "")).strip(),
                "module": require_text(row, "module", context),
                "function": require_text(row, "function", context),
                "content": require_text(row, "content", context),
                "status": require_text(row, "status", context),
                "evidence": require_text(row, "evidence", context),
                "integration": require_text(row, "integration", context),
                "pricing": require_text(row, "pricing", context),
            }
        )

    exclusions_input = data.get("exclusions")
    if not isinstance(exclusions_input, list):
        raise ValueError("exclusions 必须是数组")
    exclusions: list[str] = []
    for index, value in enumerate(exclusions_input, 1):
        text = str(value).strip() if isinstance(value, str) else ""
        if not text:
            raise ValueError(f"exclusions[{index}] 必须是非空文本")
        exclusions.append(text)

    plan_input = data.get("implementation_plan")
    if not isinstance(plan_input, list) or not plan_input:
        raise ValueError("implementation_plan 必须是非空数组")
    plan: list[dict[str, Any]] = []
    payment_ratio_total = Decimal("0")
    for index, row in enumerate(plan_input, 1):
        if not isinstance(row, dict):
            raise ValueError(f"implementation_plan[{index}] 必须是对象")
        context = f"implementation_plan[{index}]"
        ratio = parse_payment_ratio(row.get("payment_ratio"), f"{context}.payment_ratio")
        payment_ratio_total += ratio
        plan.append(
            {
                "phase": require_text(row, "phase", context),
                "period": require_text(row, "period", context),
                "work": require_text(row, "work", context),
                "deliverables": require_text(row, "deliverables", context),
                "acceptance_payment": require_text(row, "acceptance_payment", context),
                "payment_ratio": ratio,
                "payment_amount": Decimal("0"),
            }
        )
    if payment_ratio_total != Decimal("1"):
        raise ValueError(
            f"implementation_plan 付款比例合计必须为 100%，当前为 "
            f"{number_text(payment_ratio_total * 100)}%"
        )

    modules = data.get("modules")
    if not isinstance(modules, list) or not modules:
        raise ValueError("modules 必须是非空数组")

    seen_modules: set[str] = set()
    seen_items: set[str] = set()
    flat_items: list[dict[str, Any]] = []
    module_rows: list[dict[str, Any]] = []
    person_days: dict[str, Decimal] = defaultdict(lambda: Decimal("0"))
    person_amounts: dict[str, Decimal] = defaultdict(lambda: Decimal("0"))
    matrix: dict[str, dict[str, Decimal]] = defaultdict(
        lambda: defaultdict(lambda: Decimal("0"))
    )

    for module_index, module in enumerate(modules, 1):
        if not isinstance(module, dict):
            raise ValueError(f"modules[{module_index}] 必须是对象")
        module_id = require_text(module, "id", f"modules[{module_index}]")
        module_name = require_text(module, "name", f"modules[{module_index}]")
        module_context = f"modules[{module_index}]"
        module_purpose = require_text(module, "purpose", module_context)
        module_integration = require_text(module, "integration", module_context)
        module_phase = require_text(module, "phase", module_context)
        if module_id in seen_modules:
            raise ValueError(f"模块编号重复：{module_id}")
        seen_modules.add(module_id)
        items = module.get("items")
        if not isinstance(items, list) or not items:
            raise ValueError(f"模块 {module_id} 的 items 必须是非空数组")

        module_days = Decimal("0")
        module_amount = Decimal("0")
        for item_index, item in enumerate(items, 1):
            if not isinstance(item, dict):
                raise ValueError(f"{module_id}.items[{item_index}] 必须是对象")
            item_id = require_text(item, "id", f"{module_id}.items[{item_index}]")
            item_name = require_text(item, "name", item_id)
            item_description = require_text(item, "description", item_id)
            item_build_type = require_text(item, "build_type", item_id)
            item_owner_role = require_text(item, "owner_role", item_id)
            item_difficulty = require_text(item, "difficulty", item_id)
            item_phase = require_text(item, "phase", item_id)
            item_acceptance = require_text(item, "acceptance", item_id)
            item_evidence = require_text(item, "evidence_grade", item_id).upper()
            if item_evidence not in {"A", "B", "C", "D", "E"}:
                raise ValueError(f"{item_id}.evidence_grade 必须是 A 至 E")
            if item_id in seen_items:
                raise ValueError(f"功能编号重复：{item_id}")
            seen_items.add(item_id)
            days = decimal(item.get("days"), f"{item_id}.days")
            if days <= 0:
                raise ValueError(f"{item_id}.days 必须大于 0")
            if days != round_day(days):
                raise ValueError(f"{item_id}.days 最多保留两位小数")
            raw_allocations = item.get("person_days")
            if not isinstance(raw_allocations, dict) or not raw_allocations:
                raise ValueError(f"{item_id}.person_days 必须是非空对象")
            allocations: list[tuple[str, Decimal]] = []
            for name, raw_days in raw_allocations.items():
                if name not in people:
                    raise ValueError(f"{item_id} 引用了未登记人员：{name}")
                allocated_days = decimal(raw_days, f"{item_id}.person_days.{name}")
                if allocated_days < 0:
                    raise ValueError(f"{item_id} 的 {name} 人日不能小于 0")
                if allocated_days != round_day(allocated_days):
                    raise ValueError(f"{item_id} 的 {name} 人日最多保留两位小数")
                if allocated_days > 0:
                    allocations.append((name, allocated_days))
            allocated_total = sum((value for _, value in allocations), Decimal("0"))
            if allocated_total != days:
                raise ValueError(
                    f"{item_id} 人员分配 {number_text(allocated_total)} 人日，"
                    f"不等于功能工作量 {number_text(days)} 人日"
                )

            amount = round_money(days * day_rate)
            allocated_amounts = allocate_amount_by_weight(amount, allocations, days)
            for name, allocated_days in allocations:
                person_days[name] += allocated_days
                person_amounts[name] += allocated_amounts[name]
                matrix[module_id][name] += allocated_days

            row = {
                "module_id": module_id,
                "module_name": module_name,
                "module_purpose": module_purpose,
                "module_integration": module_integration,
                "module_phase": module_phase,
                "item_id": item_id,
                "item_name": item_name,
                "description": item_description,
                "build_type": item_build_type,
                "owner_role": item_owner_role,
                "difficulty": item_difficulty,
                "days": days,
                "day_rate": day_rate,
                "amount": amount,
                "phase": item_phase,
                "acceptance": item_acceptance,
                "evidence_grade": item_evidence,
                "person_days": dict(allocations),
                "person_amounts": allocated_amounts,
            }
            flat_items.append(row)
            module_days += days
            module_amount += amount

        module_rows.append(
            {
                "id": module_id,
                "name": module_name,
                "purpose": module_purpose,
                "integration": module_integration,
                "phase": module_phase,
                "days": module_days,
                "base_amount": module_amount,
                "risk_amount": round_money(module_amount * risk_rate),
                "total_amount": Decimal("0"),
            }
        )

    total_days = sum((row["days"] for row in module_rows), Decimal("0"))
    base_amount = sum((row["base_amount"] for row in module_rows), Decimal("0"))
    risk_amount = round_money(base_amount * risk_rate)
    module_risk = allocate_amount_by_weight(
        risk_amount,
        [(row["id"], row["base_amount"]) for row in module_rows],
        base_amount,
    )
    for row in module_rows:
        row["risk_amount"] = module_risk[row["id"]]
        row["total_amount"] = row["base_amount"] + row["risk_amount"]

    total_amount = base_amount + risk_amount
    payment_amounts = allocate_amount_by_weight(
        total_amount,
        [(str(index), row["payment_ratio"]) for index, row in enumerate(plan)],
        Decimal("1"),
    )
    for index, row in enumerate(plan):
        row["payment_amount"] = payment_amounts[str(index)]

    if sum(person_days.values(), Decimal("0")) != total_days:
        raise AssertionError("人员总人日与项目总人日不一致")
    if sum(person_amounts.values(), Decimal("0")) != base_amount:
        raise AssertionError("人员分摊金额与建设基价不一致")
    if any(value < 0 for value in person_amounts.values()):
        raise AssertionError("人员分摊金额不能为负数")
    for row in module_rows:
        matrix_total = sum(matrix[row["id"]].values(), Decimal("0"))
        if matrix_total != row["days"]:
            raise AssertionError(f"模块 {row['id']} 的人员矩阵与模块人日不一致")
        if row["risk_amount"] < 0 or row["total_amount"] < 0:
            raise AssertionError(f"模块 {row['id']} 的风险或预算金额不能为负数")
    if sum((row["payment_amount"] for row in plan), Decimal("0")) != total_amount:
        raise AssertionError("阶段付款金额与预算总价不一致")

    person_rows: list[dict[str, Any]] = []
    for name, person in people.items():
        days = person_days[name]
        person_rows.append(
            {
                "name": name,
                "role": str(person.get("role", "")).strip(),
                "responsibility": str(person.get("responsibility", "")).strip(),
                "days": days,
                "person_months": days / month_days,
                "share": days / total_days if total_days else Decimal("0"),
                "base_amount": person_amounts[name],
            }
        )

    return {
        "source": data,
        "project_name": project_name,
        "meta": meta,
        "day_rate": day_rate,
        "risk_rate": risk_rate,
        "total_days": total_days,
        "person_months": total_days / month_days,
        "base_amount": base_amount,
        "risk_amount": risk_amount,
        "total_amount": total_amount,
        "items": flat_items,
        "modules": module_rows,
        "people": person_rows,
        "matrix": matrix,
        "current_system_baseline": baseline,
        "implementation_plan": plan,
        "exclusions": exclusions,
    }


def serializable(result: dict[str, Any]) -> dict[str, Any]:
    money_fields = {
        "day_rate", "amount", "base_amount", "risk_amount", "total_amount", "payment_amount"
    }
    day_fields = {"days", "total_days"}
    context_fields = {"person_amounts", "person_days", "matrix"}

    def convert(value: Any, key: str = "", context: str = "") -> Any:
        if isinstance(value, Decimal):
            if key in money_fields or context == "person_amounts":
                return money_text(value)
            if key in day_fields or context in {"person_days", "matrix"}:
                return money_text(value)
            return number_text(value)
        if isinstance(value, defaultdict):
            value = dict(value)
        if isinstance(value, dict):
            next_context = key if key in context_fields else context
            return {
                str(item_key): convert(item, str(item_key), next_context)
                for item_key, item in value.items()
            }
        if isinstance(value, list):
            return [convert(item, key, context) for item in value]
        return value

    return convert({key: value for key, value in result.items() if key != "source"})


def allocation_text(row: dict[str, Any]) -> str:
    return "、".join(
        f"{name}{number_text(days)}人日" for name, days in row["person_days"].items()
    )


def write_json(result: dict[str, Any], path: Path) -> None:
    path.write_text(
        json.dumps(serializable(result), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def write_csv(result: dict[str, Any], path: Path) -> None:
    fields = [
        "模块编号", "模块名称", "功能编号", "功能名称", "主要内容", "建设属性", "难度",
        "人日", "人日单价（元）", "合价（元）", "人员分配", "验收标准", "阶段", "证据等级",
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        for row in result["items"]:
            writer.writerow(
                {
                    "模块编号": spreadsheet_safe(row["module_id"]),
                    "模块名称": spreadsheet_safe(row["module_name"]),
                    "功能编号": spreadsheet_safe(row["item_id"]),
                    "功能名称": spreadsheet_safe(row["item_name"]),
                    "主要内容": spreadsheet_safe(row["description"]),
                    "建设属性": spreadsheet_safe(row["build_type"]),
                    "难度": spreadsheet_safe(row["difficulty"]),
                    "人日": number_text(row["days"]),
                    "人日单价（元）": money_text(row["day_rate"]),
                    "合价（元）": money_text(row["amount"]),
                    "人员分配": spreadsheet_safe(allocation_text(row)),
                    "验收标准": spreadsheet_safe(row["acceptance"]),
                    "阶段": spreadsheet_safe(row["phase"]),
                    "证据等级": spreadsheet_safe(row["evidence_grade"]),
                }
            )


def markdown_table(headers: list[str], rows: Iterable[Iterable[str]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "|" + "|".join(["---"] * len(headers)) + "|"]
    for row in rows:
        lines.append("| " + " | ".join(str(value).replace("|", "\\|").replace("\n", " ") for value in row) + " |")
    return "\n".join(lines)


def write_markdown(result: dict[str, Any], path: Path) -> None:
    meta = result["meta"]
    lines = [
        f"# {result['project_name']}工作量统计与报价方案",
        "",
        "## 一、编制口径",
        "",
        f"- 价格基准日：{meta.get('price_base_date', '未提供')}。",
        f"- 计价口径：{meta.get('tax_basis', '未提供')}；币种：{meta.get('currency', 'CNY')}。",
        f"- 月综合费率：{money_text(decimal(meta['monthly_rate_cny'], 'monthly_rate_cny'))} 元/人月；"
        f"每人月：{number_text(decimal(meta['person_month_days'], 'person_month_days'))} 人日；"
        f"折算人日单价：{money_text(result['day_rate'])} 元。",
        f"- 风险准备率：{number_text(result['risk_rate'] * 100)}%。",
        "- 公式：功能合价=ROUND(功能人日×人日单价,2)；风险准备=ROUND(建设基价×风险准备率,2)。",
        "",
        "## 二、报价汇总",
        "",
        markdown_table(
            ["项目", "结果"],
            [
                ["总工作量", f"{number_text(result['total_days'])} 人日"],
                ["折算工作量", f"{result['person_months']:.2f} 人月"],
                ["建设基价", f"{money_text(result['base_amount'])} 元"],
                ["风险准备", f"{money_text(result['risk_amount'])} 元"],
                ["预算总价", f"{money_text(result['total_amount'])} 元"],
            ],
        ),
        "",
        "## 三、模块汇总",
        "",
        markdown_table(
            ["编号", "模块", "主要用途", "人日", "建设基价", "风险准备", "合计"],
            ([
                row["id"], row["name"], row["purpose"], number_text(row["days"]),
                money_text(row["base_amount"]), money_text(row["risk_amount"]), money_text(row["total_amount"]),
            ] for row in result["modules"]),
        ),
        "",
        "## 四、分部分项工作量与报价",
        "",
        markdown_table(
            ["功能编号", "模块", "功能", "主要内容", "属性", "难度", "阶段", "证据", "人日", "人日单价", "合价", "人员分配", "验收标准"],
            ([
                row["item_id"], row["module_name"], row["item_name"], row["description"],
                row["build_type"], row["difficulty"], row["phase"], row["evidence_grade"],
                number_text(row["days"]), money_text(row["day_rate"]), money_text(row["amount"]),
                allocation_text(row), row["acceptance"],
            ] for row in result["items"]),
        ),
        "",
        "## 五、人员职责与工作量",
        "",
        markdown_table(
            ["人员", "角色", "职责", "人日", "人月", "占比", "分摊建设基价"],
            ([
                row["name"], row["role"], row["responsibility"], number_text(row["days"]),
                f"{row['person_months']:.2f}", f"{row['share'] * 100:.2f}%", money_text(row["base_amount"]),
            ] for row in result["people"]),
        ),
        "",
        "## 六、现有系统功能基线",
        "",
        markdown_table(
            ["编号", "现有模块", "功能", "主要内容", "状态", "证据", "融合处理", "计价处理"],
            ([row["id"], row["module"], row["function"], row["content"], row["status"], row["evidence"], row["integration"], row["pricing"]] for row in result["current_system_baseline"]),
        ),
        "",
        "## 七、模块—人员工作量矩阵",
        "",
        markdown_table(
            ["模块编号", "模块"] + [row["name"] for row in result["people"]] + ["模块合计"],
            ([module["id"], module["name"]] + [number_text(result["matrix"][module["id"]][person["name"]]) for person in result["people"]] + [number_text(module["days"])] for module in result["modules"]),
        ),
        "",
        "## 八、实施与付款计划",
        "",
        markdown_table(
            ["阶段", "时间", "主要工作", "交付物", "验收/付款条件", "比例", "付款金额"],
            ([row["phase"], row["period"], row["work"], row["deliverables"], row["acceptance_payment"], f"{number_text(row['payment_ratio'] * 100)}%", money_text(row["payment_amount"])] for row in result["implementation_plan"]),
        ),
        "",
        "## 九、范围与限制",
        "",
    ]
    exclusions = result["exclusions"]
    if exclusions:
        lines.extend(f"- {item}" for item in exclusions)
    else:
        lines.append("- 未提供排除项，正式报价前应补充范围边界。")
    lines.extend([
        "- 现有系统功能清单用于确定复用、改造和回归范围；未列入 WBS 的既有功能不重复计价。",
        "- 人员人日、模块人日、功能人日和建设基价已执行勾稽校验。",
        "",
    ])
    path.write_text("\n".join(lines), encoding="utf-8")


def _excel_imports():
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise RuntimeError("生成 XLSX 需要 openpyxl：python3 -m pip install -r requirements-docs.txt") from exc
    return Workbook, Alignment, Border, Font, PatternFill, Side, get_column_letter


def write_xlsx(result: dict[str, Any], path: Path) -> None:
    Workbook, Alignment, Border, Font, PatternFill, Side, get_column_letter = _excel_imports()
    workbook = Workbook()
    workbook.remove(workbook.active)
    navy, blue, pale, white = "17365D", "5B9BD5", "D9EAF7", "FFFFFF"
    thin = Side(style="thin", color="B7C9D6")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def configure(sheet, widths: list[float]) -> None:
        sheet.sheet_view.showGridLines = False
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        sheet.page_setup.paperSize = sheet.PAPERSIZE_A4
        sheet.page_setup.orientation = sheet.ORIENTATION_PORTRAIT
        sheet.page_setup.fitToWidth = 1
        sheet.sheet_properties.pageSetUpPr.fitToPage = True
        sheet.page_margins.left = 0.2
        sheet.page_margins.right = 0.2
        sheet.page_margins.top = 0.35
        sheet.page_margins.bottom = 0.35
        for index, width in enumerate(widths, 1):
            sheet.column_dimensions[get_column_letter(index)].width = width
        for cell in sheet[1]:
            cell.fill = PatternFill("solid", fgColor=navy)
            cell.font = Font(color=white, bold=True, size=9)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
        for row in sheet.iter_rows(min_row=2):
            for cell in row:
                cell.font = Font(size=8)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = border

    details = workbook.create_sheet("分部分项清单")
    details.append(["模块编号", "模块", "功能编号", "功能", "主要内容", "属性", "难度", "人日", "人日单价", "合价", "人员分配", "验收标准"])
    for row in result["items"]:
        details.append([
            spreadsheet_safe(row["module_id"]), spreadsheet_safe(row["module_name"]),
            spreadsheet_safe(row["item_id"]), spreadsheet_safe(row["item_name"]),
            spreadsheet_safe(row["description"]), spreadsheet_safe(row["build_type"]),
            spreadsheet_safe(row["difficulty"]), float(row["days"]), float(row["day_rate"]),
            f"=ROUND(H{details.max_row + 1}*I{details.max_row + 1},2)",
            spreadsheet_safe(allocation_text(row)), spreadsheet_safe(row["acceptance"]),
        ])
    configure(details, [10, 16, 11, 18, 34, 10, 8, 8, 12, 12, 24, 30])
    details.auto_filter.ref = f"A1:L{details.max_row}"
    for row in range(2, details.max_row + 1):
        details.cell(row, 8).number_format = "0.00"
        details.cell(row, 9).number_format = '#,##0.00" 元"'
        details.cell(row, 10).number_format = '#,##0.00" 元"'

    summary = workbook.create_sheet("模块汇总")
    summary.append(["编号", "模块", "主要用途", "与现有系统融合", "人日", "建设基价", "风险准备", "预算合计"])
    for row in result["modules"]:
        summary.append([spreadsheet_safe(row["id"]), spreadsheet_safe(row["name"]), spreadsheet_safe(row["purpose"]), spreadsheet_safe(row["integration"]), float(row["days"]), float(row["base_amount"]), float(row["risk_amount"]), float(row["total_amount"])])
    summary.append(["", "合计", "", "", float(result["total_days"]), float(result["base_amount"]), float(result["risk_amount"]), float(result["total_amount"])])
    configure(summary, [9, 18, 34, 32, 9, 14, 13, 14])
    for cell in summary[summary.max_row]:
        cell.fill = PatternFill("solid", fgColor=pale)
        cell.font = Font(bold=True, size=8)
    for row in range(2, summary.max_row + 1):
        for column in range(6, 9):
            summary.cell(row, column).number_format = '#,##0.00" 元"'

    baseline = workbook.create_sheet("现有系统基线")
    baseline.append(["编号", "现有模块", "功能", "主要内容", "当前状态", "证据", "融合处理", "计价处理"])
    for row in result["current_system_baseline"]:
        baseline.append([spreadsheet_safe(row["id"]), spreadsheet_safe(row["module"]), spreadsheet_safe(row["function"]), spreadsheet_safe(row["content"]), spreadsheet_safe(row["status"]), spreadsheet_safe(row["evidence"]), spreadsheet_safe(row["integration"]), spreadsheet_safe(row["pricing"])])
    configure(baseline, [10, 18, 18, 32, 12, 22, 28, 16])

    people_sheet = workbook.create_sheet("人员工作量")
    people_sheet.append(["人员", "角色", "职责", "人日", "折算人月", "工作量占比", "分摊建设基价"])
    for row in result["people"]:
        people_sheet.append([spreadsheet_safe(row["name"]), spreadsheet_safe(row["role"]), spreadsheet_safe(row["responsibility"]), float(row["days"]), float(row["person_months"]), float(row["share"]), float(row["base_amount"])])
    configure(people_sheet, [14, 16, 38, 10, 11, 12, 16])
    for row in range(2, people_sheet.max_row + 1):
        people_sheet.cell(row, 6).number_format = "0.00%"
        people_sheet.cell(row, 7).number_format = '#,##0.00" 元"'

    matrix_sheet = workbook.create_sheet("模块人员矩阵")
    names = [row["name"] for row in result["people"]]
    matrix_sheet.append(["模块编号", "模块"] + [spreadsheet_safe(name) for name in names] + ["模块合计"])
    module_map = {row["id"]: row for row in result["modules"]}
    for module_id, row in module_map.items():
        values = [float(result["matrix"][module_id][name]) for name in names]
        matrix_sheet.append([spreadsheet_safe(module_id), spreadsheet_safe(row["name"])] + values + [float(row["days"])])
    configure(matrix_sheet, [11, 20] + [12] * len(names) + [12])

    plan_sheet = workbook.create_sheet("实施与付款计划")
    plan_sheet.append(["阶段", "时间", "主要工作", "交付物", "验收/付款条件", "建议比例", "建议付款金额"])
    for row in result["implementation_plan"]:
        plan_sheet.append([spreadsheet_safe(row["phase"]), spreadsheet_safe(row["period"]), spreadsheet_safe(row["work"]), spreadsheet_safe(row["deliverables"]), spreadsheet_safe(row["acceptance_payment"]), float(row["payment_ratio"]), float(row["payment_amount"])])
    configure(plan_sheet, [14, 14, 31, 28, 31, 12, 16])
    for row in range(2, plan_sheet.max_row + 1):
        plan_sheet.cell(row, 6).number_format = "0.00%"
        plan_sheet.cell(row, 7).number_format = '#,##0.00" 元"'

    notes = workbook.create_sheet("编制说明")
    notes.append(["项目", "内容"])
    notes_rows = [
        ("项目名称", spreadsheet_safe(result["project_name"])),
        ("价格基准日", spreadsheet_safe(result["meta"].get("price_base_date", "未提供"))),
        ("计价口径", spreadsheet_safe(result["meta"].get("tax_basis", "未提供"))),
        ("人日单价", f"月综合费率÷每人月人日={money_text(result['day_rate'])} 元/人日"),
        ("总价公式", "各功能 ROUND(人日×人日单价,2) 后求和，再加风险准备"),
        ("勾稽关系", "功能人日=人员分配人日；模块人日=模块人员矩阵；个人金额合计=建设基价"),
        ("范围规则", "现有系统基线用于复用、融合和回归识别；未列入 WBS 的既有功能不重复计价"),
    ]
    for row in notes_rows:
        notes.append(row)
    for exclusion in result["exclusions"]:
        notes.append(["排除项", spreadsheet_safe(exclusion)])
    configure(notes, [18, 85])

    workbook.save(path)


def _docx_imports():
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise RuntimeError("生成 DOCX 需要 python-docx：python3 -m pip install -r requirements-docs.txt") from exc
    return Document, WD_ORIENT, WD_ALIGN_PARAGRAPH, Cm, Pt


def write_docx(result: dict[str, Any], path: Path) -> None:
    Document, WD_ORIENT, WD_ALIGN_PARAGRAPH, Cm, Pt = _docx_imports()
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(1.3)
    section.left_margin = section.right_margin = Cm(1.2)
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{result['project_name']}\n工作量统计与报价方案")
    run.bold = True
    run.font.size = Pt(16)

    def add_table(headers: list[str], rows: Iterable[Iterable[Any]], widths: list[float] | None = None):
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = False
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = header
            if widths:
                cell.width = Cm(widths[index])
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for item in paragraph.runs:
                    item.bold = True
                    item.font.size = Pt(8)
        for values in rows:
            cells = table.add_row().cells
            for index, value in enumerate(values):
                cells[index].text = str(value)
                if widths:
                    cells[index].width = Cm(widths[index])
                for paragraph in cells[index].paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1
                    for item in paragraph.runs:
                        item.font.size = Pt(7.5)
        return table

    document.add_heading("一、编制与计价说明", level=1)
    meta = result["meta"]
    document.add_paragraph(
        f"价格基准日：{meta.get('price_base_date', '未提供')}；计价口径：{meta.get('tax_basis', '未提供')}。"
        f"月综合费率 {money_text(decimal(meta['monthly_rate_cny'], 'monthly_rate_cny'))} 元，"
        f"按 {number_text(decimal(meta['person_month_days'], 'person_month_days'))} 人日/人月折算，"
        f"人日单价 {money_text(result['day_rate'])} 元。"
    )
    add_table(
        ["总工作量", "折算人月", "建设基价", "风险准备", "预算总价"],
        [[f"{number_text(result['total_days'])} 人日", f"{result['person_months']:.2f}", f"{money_text(result['base_amount'])} 元", f"{money_text(result['risk_amount'])} 元", f"{money_text(result['total_amount'])} 元"]],
        [3.2, 3.0, 4.0, 3.6, 4.0],
    )

    document.add_heading("二、模块工作量与报价汇总", level=1)
    add_table(
        ["编号", "模块", "主要用途", "人日", "建设基价", "风险准备", "合计"],
        ([row["id"], row["name"], row["purpose"], number_text(row["days"]), money_text(row["base_amount"]), money_text(row["risk_amount"]), money_text(row["total_amount"])] for row in result["modules"]),
        [1.5, 3.2, 5.0, 1.5, 2.5, 2.3, 2.5],
    )

    document.add_heading("三、分部分项工程量清单与计价", level=1)
    for module in result["modules"]:
        document.add_heading(f"{module['id']} {module['name']}", level=2)
        if module["purpose"]:
            document.add_paragraph(f"主要用途：{module['purpose']}")
        module_items = [row for row in result["items"] if row["module_id"] == module["id"]]
        add_table(
            ["编号", "功能/属性", "主要内容", "人日", "单价", "合价", "人员分配", "验收标准"],
            ([row["item_id"], f"{row['item_name']}\n{row['build_type']}｜{row['difficulty']}｜{row['phase']}｜证据{row['evidence_grade']}", row["description"], number_text(row["days"]), money_text(row["day_rate"]), money_text(row["amount"]), allocation_text(row), row["acceptance"]] for row in module_items),
            [1.6, 2.5, 4.2, 1.1, 1.7, 1.8, 3.0, 3.8],
        )

    document.add_heading("四、人员职责与工作量统计", level=1)
    add_table(
        ["人员", "角色", "主要职责", "人日", "人月", "占比", "分摊建设基价"],
        ([row["name"], row["role"], row["responsibility"], number_text(row["days"]), f"{row['person_months']:.2f}", f"{row['share'] * 100:.2f}%", money_text(row["base_amount"])] for row in result["people"]),
        [2.0, 2.3, 6.5, 1.4, 1.4, 1.6, 3.0],
    )

    document.add_heading("五、现有系统功能基线", level=1)
    baseline = result["current_system_baseline"]
    if baseline:
        add_table(
            ["编号", "模块/功能", "内容/状态", "证据", "融合处理", "计价处理"],
            ([row["id"], f"{row['module']}\n{row['function']}", f"{row['content']}\n状态：{row['status']}", row["evidence"], row["integration"], row["pricing"]] for row in baseline),
            [1.6, 3.0, 4.0, 2.8, 4.2, 3.5],
        )

    document.add_heading("六、模块—人员工作量矩阵", level=1)
    names = [row["name"] for row in result["people"]]
    add_table(
        ["模块编号", "模块"] + names + ["模块合计"],
        ([module["id"], module["name"]] + [number_text(result["matrix"][module["id"]][name]) for name in names] + [number_text(module["days"])] for module in result["modules"]),
        [1.7, 3.5] + [2.3] * len(names) + [2.2],
    )

    document.add_heading("七、实施与付款计划", level=1)
    add_table(
        ["阶段", "时间", "主要工作", "交付物", "验收/付款条件", "比例", "付款金额"],
        ([row["phase"], row["period"], row["work"], row["deliverables"], row["acceptance_payment"], f"{number_text(row['payment_ratio'] * 100)}%", money_text(row["payment_amount"])] for row in result["implementation_plan"]),
        [1.8, 2.0, 3.7, 3.5, 4.3, 1.4, 2.4],
    )

    document.add_heading("八、范围边界与勾稽说明", level=1)
    for exclusion in result["exclusions"]:
        document.add_paragraph(exclusion, style="List Bullet")
    document.add_paragraph("勾稽校验：功能人日等于人员分配人日；模块人日等于模块人员矩阵；人员分摊金额合计等于建设基价。")
    document.save(path)


WRITERS = {
    "json": ("-calculation.json", write_json),
    "csv": ("-details.csv", write_csv),
    "md": ("-quotation.md", write_markdown),
    "xlsx": ("-quotation.xlsx", write_xlsx),
    "docx": ("-quotation.docx", write_docx),
}


def parse_formats(raw: str) -> list[str]:
    if raw.strip().lower() == "all":
        return list(WRITERS)
    formats = [item.strip().lower() for item in raw.split(",") if item.strip()]
    unknown = sorted(set(formats) - set(WRITERS))
    if unknown:
        raise ValueError(f"不支持的格式：{', '.join(unknown)}")
    if not formats:
        raise ValueError("至少选择一种输出格式")
    return formats


def preflight_formats(formats: list[str]) -> None:
    """写盘前完成可选依赖检查，避免生成半套成果。"""
    if "xlsx" in formats:
        _excel_imports()
    if "docx" in formats:
        _docx_imports()


def publish_staged(staged: list[tuple[Path, Path]], temporary_dir: Path) -> list[Path]:
    """发布整套成果；任一替换失败时恢复发布前状态。"""
    for _, final_path in staged:
        if final_path.is_dir():
            raise ValueError(f"输出目标不能是目录：{final_path}")

    backup_dir = temporary_dir / "backups"
    backup_dir.mkdir()
    backups: dict[Path, Path] = {}
    installed: list[Path] = []
    try:
        for index, (_, final_path) in enumerate(staged):
            if os.path.lexists(final_path):
                backup_path = backup_dir / f"{index}-{final_path.name}"
                os.replace(final_path, backup_path)
                backups[final_path] = backup_path
        for temporary_path, final_path in staged:
            os.replace(temporary_path, final_path)
            installed.append(final_path)
    except Exception as exc:
        rollback_errors: list[str] = []
        restored: set[Path] = set()
        for final_path in reversed(installed):
            try:
                if final_path in backups:
                    os.replace(backups[final_path], final_path)
                    restored.add(final_path)
                elif os.path.lexists(final_path):
                    final_path.unlink()
            except Exception as rollback_exc:  # pragma: no cover - 仅文件系统二次故障
                rollback_errors.append(f"{final_path}: {rollback_exc}")
        for final_path, backup_path in backups.items():
            if final_path in restored or not os.path.lexists(backup_path):
                continue
            try:
                os.replace(backup_path, final_path)
            except Exception as rollback_exc:  # pragma: no cover - 仅文件系统二次故障
                rollback_errors.append(f"{final_path}: {rollback_exc}")
        detail = f"；回滚失败：{'；'.join(rollback_errors)}" if rollback_errors else ""
        raise RuntimeError(f"成果发布失败，已回滚：{exc}{detail}") from exc
    return [final_path for _, final_path in staged]


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="生成工作量统计与报价方案")
    parser.add_argument("input", type=Path, help="详细报价输入 JSON")
    parser.add_argument("--output-dir", type=Path, default=Path.cwd(), help="输出目录")
    parser.add_argument("--prefix", help="输出文件名前缀，默认使用输入文件名")
    parser.add_argument("--formats", default="all", help="json,csv,md,xlsx,docx 或 all")
    args = parser.parse_args(argv)
    try:
        data = json.loads(
            args.input.read_text(encoding="utf-8"),
            parse_constant=lambda value: (_ for _ in ()).throw(
                ValueError(f"JSON 不允许非有限数：{value}")
            ),
        )
        result = prepare(data)
        formats = parse_formats(args.formats)
        preflight_formats(formats)
        args.output_dir.mkdir(parents=True, exist_ok=True)
        prefix = args.prefix or args.input.stem
        if Path(prefix).name != prefix or prefix in {"", ".", ".."}:
            raise ValueError("输出前缀必须是单个有效文件名，不能包含路径")
        output_paths: list[Path] = []
        with tempfile.TemporaryDirectory(prefix=".quotation-", dir=args.output_dir) as directory:
            temporary_dir = Path(directory)
            staged: list[tuple[Path, Path]] = []
            for output_format in formats:
                suffix, writer = WRITERS[output_format]
                final_path = args.output_dir / f"{prefix}{suffix}"
                temporary_path = temporary_dir / final_path.name
                writer(result, temporary_path)
                staged.append((temporary_path, final_path))
            output_paths = publish_staged(staged, temporary_dir)
    except (OSError, ValueError, RuntimeError, json.JSONDecodeError) as exc:
        print(f"错误：{exc}", file=sys.stderr)
        return 2
    print(
        f"已生成：{', '.join(str(path) for path in output_paths)}\n"
        f"总工作量 {number_text(result['total_days'])} 人日，预算总价 {money_text(result['total_amount'])} 元。"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
